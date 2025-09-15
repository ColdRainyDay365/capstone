from odoo import models, fields, api
from odoo.exceptions import UserError
import base64
import os
from docxtpl import DocxTemplate
from io import BytesIO

class ChangeRequestExportWizard(models.TransientModel):
    _name = 'change.request.export.wizard'
    _description = 'Change Request Export Wizard'
    
    change_request_id = fields.Many2one(
        'simple.change.request',
        string="Change Request",
        required=True
    )
    
    template_file = fields.Binary(
        string="Word Template",
        help="Upload a Word template (.docx) file for mail merge"
    )
    
    template_filename = fields.Char(
        string="Template Filename"
    )
    
    output_filename = fields.Char(
        string="Output Filename",
        default="Change_Request_{}.docx"
    )
    
    include_effort_breakdown = fields.Boolean(
        string="Include Effort Breakdown",
        default=True,
        help="Include the effort breakdown table in the export"
    )
    

    
    @api.model
    def default_get(self, fields_list):
        """Set default values"""
        res = super().default_get(fields_list)
        if self.env.context.get('default_change_request_id'):
            change_request = self.env['simple.change.request'].browse(
                self.env.context.get('default_change_request_id')
            )
            res['output_filename'] = f"Change_Request_{change_request.change_number or change_request.name}.docx"
        return res
    
    def action_export_to_word(self):
        """Export change request to Word document"""
        self.ensure_one()
        
        if not self.template_file:
            # Create a basic Word document without template
            return self._create_basic_word_document()
        else:
            # Use the uploaded template for mail merge
            return self._create_template_word_document()
    
    def _create_basic_word_document(self):
        """Create a basic Word document without template"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Create a new document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Change Request', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add basic information
            doc.add_heading('Basic Information', level=1)
            
            # Create a table for basic info
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Field'
            hdr_cells[1].text = 'Value'
            
            # Add basic information rows
            basic_info = [
                ('Project Name', self.change_request_id.project_name or ''),
                ('Change Number', self.change_request_id.change_number or ''),
                ('Change Request Date', str(self.change_request_id.change_request_date) if self.change_request_id.change_request_date else ''),
                ('Requester', self.change_request_id.requester_id.name if self.change_request_id.requester_id else ''),
                ('Department', self.change_request_id.department or ''),
                ('Request Type', dict(self.change_request_id._fields['request_type'].selection).get(self.change_request_id.request_type, '')),
                ('Priority', dict(self.change_request_id._fields['priority'].selection).get(self.change_request_id.priority, '')),
                ('Expected Completion', str(self.change_request_id.expected_completion) if self.change_request_id.expected_completion else ''),
            ]
            
            for field, value in basic_info:
                row_cells = table.add_row().cells
                row_cells[0].text = field
                row_cells[1].text = value
            
            # Add problem statement
            if self.change_request_id.problem_statement:
                doc.add_heading('Problem Statement', level=1)
                doc.add_paragraph(self.change_request_id.problem_statement)
            
            # Add change description
            if self.change_request_id.change_description:
                doc.add_heading('Change Description', level=1)
                # Convert HTML to plain text for basic document
                import re
                clean_text = re.sub('<[^<]+?>', '', self.change_request_id.change_description)
                doc.add_paragraph(clean_text)
            
            # Add acceptance criteria
            if self.change_request_id.acceptance_criteria:
                doc.add_heading('Acceptance Criteria', level=1)
                doc.add_paragraph(self.change_request_id.acceptance_criteria)
            
            # Add reason and benefits
            if self.change_request_id.reason_for_change or self.change_request_id.benefits_of_change:
                doc.add_heading('Reason and Benefits', level=1)
                if self.change_request_id.reason_for_change:
                    doc.add_heading('Reason for Change', level=2)
                    doc.add_paragraph(self.change_request_id.reason_for_change)
                if self.change_request_id.benefits_of_change:
                    doc.add_heading('Benefits of Change', level=2)
                    doc.add_paragraph(self.change_request_id.benefits_of_change)
            
            # Add schedule and cost
            if self.change_request_id.delivery_timeline or self.change_request_id.cost_estimation:
                doc.add_heading('Schedule and Cost', level=1)
                if self.change_request_id.delivery_timeline:
                    doc.add_heading('Delivery Timeline', level=2)
                    doc.add_paragraph(self.change_request_id.delivery_timeline)
                if self.change_request_id.cost_estimation:
                    doc.add_heading('Cost Estimation', level=2)
                    doc.add_paragraph(self.change_request_id.cost_estimation)
            
            # Add assumptions and payment
            if self.change_request_id.assumptions or self.change_request_id.payment_milestones:
                doc.add_heading('Assumptions and Payment', level=1)
                if self.change_request_id.assumptions:
                    doc.add_heading('Assumptions', level=2)
                    doc.add_paragraph(self.change_request_id.assumptions)
                if self.change_request_id.payment_milestones:
                    doc.add_heading('Payment Milestones', level=2)
                    doc.add_paragraph(self.change_request_id.payment_milestones)
            
            # Add effort breakdown if requested
            if self.include_effort_breakdown and self.change_request_id.effort_breakdown_ids:
                doc.add_heading('Effort Breakdown', level=1)
                
                # Create effort breakdown table
                effort_table = doc.add_table(rows=1, cols=4)
                effort_table.style = 'Table Grid'
                effort_hdr_cells = effort_table.rows[0].cells
                effort_hdr_cells[0].text = 'No.'
                effort_hdr_cells[1].text = 'Expected Task'
                effort_hdr_cells[2].text = 'Effort (Days)'
                effort_hdr_cells[3].text = 'Remarks'
                
                for effort_line in self.change_request_id.effort_breakdown_ids:
                    row_cells = effort_table.add_row().cells
                    row_cells[0].text = effort_line.task_number or ''
                    row_cells[1].text = effort_line.expected_task or ''
                    row_cells[2].text = str(effort_line.effort_days) if effort_line.effort_days else ''
                    row_cells[3].text = effort_line.remarks or ''
                
                # Add total effort
                doc.add_paragraph(f'Total Effort: {self.change_request_id.total_effort_days} days')
            
            # Save the document
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            
            # Create attachment
            filename = self.output_filename.format(self.change_request_id.change_number or self.change_request_id.name)
            attachment = self.env['ir.attachment'].create({
                'name': filename,
                'type': 'binary',
                'datas': base64.b64encode(output.getvalue()),
                'res_model': 'simple.change.request',
                'res_id': self.change_request_id.id,
            })
            
            return {
                'type': 'ir.actions.act_url',
                'url': f'/web/content/{attachment.id}?download=true',
                'target': 'self',
            }
            
        except ImportError:
            raise UserError("Please install python-docx library: pip install python-docx")
        except Exception as e:
            raise UserError(f"Error creating Word document: {str(e)}")
    
    def _create_template_word_document(self):
        """Create Word document using uploaded template"""
        try:
            # Decode template file
            template_data = base64.b64decode(self.template_file)
            template_stream = BytesIO(template_data)
            
            # Load template
            doc = DocxTemplate(template_stream)
            
            # Prepare context data
            context = self._prepare_template_context()
            
            # Render template
            doc.render(context)
            
            # Save the document
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            
            # Create attachment
            filename = self.output_filename.format(self.change_request_id.change_number or self.change_request_id.name)
            attachment = self.env['ir.attachment'].create({
                'name': filename,
                'type': 'binary',
                'datas': base64.b64encode(output.getvalue()),
                'res_model': 'simple.change.request',
                'res_id': self.change_request_id.id,
            })
            
            return {
                'type': 'ir.actions.act_url',
                'url': f'/web/content/{attachment.id}?download=true',
                'target': 'self',
            }
            
        except ImportError:
            raise UserError("Please install docxtpl library: pip install docxtpl")
        except Exception as e:
            raise UserError(f"Error creating Word document from template: {str(e)}")
    
    def _prepare_template_context(self):
        """Prepare context data for template rendering"""
        cr = self.change_request_id
        
        # Basic information
        context = {
            'project_name': cr.project_name or '',
            'change_number': cr.change_number or '',
            'change_request_date': str(cr.change_request_date) if cr.change_request_date else '',
            'requester_name': cr.requester_id.name if cr.requester_id else '',
            'department': cr.department or '',
            'request_type': dict(cr._fields['request_type'].selection).get(cr.request_type, ''),
            'priority': dict(cr._fields['priority'].selection).get(cr.priority, ''),
            'expected_completion': str(cr.expected_completion) if cr.expected_completion else '',
            'problem_statement': cr.problem_statement or '',
            'change_description': cr.change_description or '',
            'acceptance_criteria': cr.acceptance_criteria or '',
            'reason_for_change': cr.reason_for_change or '',
            'benefits_of_change': cr.benefits_of_change or '',
            'delivery_timeline': cr.delivery_timeline or '',
            'cost_estimation': cr.cost_estimation or '',
            'assumptions': cr.assumptions or '',
            'payment_milestones': cr.payment_milestones or '',
            'total_effort_days': cr.total_effort_days or 0,
        }
        
        # Effort breakdown table
        if self.include_effort_breakdown and cr.effort_breakdown_ids:
            effort_breakdown = []
            for line in cr.effort_breakdown_ids:
                effort_breakdown.append({
                    'task_number': line.task_number or '',
                    'expected_task': line.expected_task or '',
                    'effort_days': line.effort_days or 0,
                    'remarks': line.remarks or '',
                    'task_category': dict(line._fields['task_category'].selection).get(line.task_category, ''),
                })
            context['effort_breakdown'] = effort_breakdown
        
        return context
